from PyPDF2 import PdfReader
from docx import Document
import os

def get_file_extension(filename):
    """Dosya uzantısını döndürür."""
    return os.path.splitext(filename)[1].lower()

def get_pdf_text(pdf_file):
    """Tek bir PDF dosyasından metin çıkarır."""
    text = ""
    try:
        pdf_reader = PdfReader(pdf_file)
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"PDF okuma hatası: {e}")
    return text.strip()

def get_docx_text(docx_file):
    """Tek bir DOCX dosyasından metin çıkarır."""
    text = ""
    try:
        doc = Document(docx_file)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Tablolardaki metinleri de al
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        print(f"DOCX okuma hatası: {e}")
    return text.strip()

def get_document_text(uploaded_files):
    """
    Yüklenen dosyalardan metin çıkarır.
    Desteklenen formatlar: PDF, DOCX
    
    Returns:
        list: Her dosya için (filename, text, doc_type) tuple'ları
    """
    documents = []
    
    for file in uploaded_files:
        filename = file.name
        extension = get_file_extension(filename)
        
        if extension == '.pdf':
            text = get_pdf_text(file)
            doc_type = 'pdf'
        elif extension in ['.docx', '.doc']:
            text = get_docx_text(file)
            doc_type = 'docx'
        else:
            print(f"Desteklenmeyen dosya formatı: {extension}")
            continue
        
        if text:
            documents.append({
                'filename': filename,
                'content': text,
                'doc_type': doc_type
            })
    
    return documents

def get_combined_text(uploaded_files):
    """
    Tüm yüklenen dosyalardan tek bir metin oluşturur.
    (Mevcut pdf_handler.py ile uyumluluk için)
    """
    documents = get_document_text(uploaded_files)
    combined = ""
    for doc in documents:
        combined += f"\n\n--- {doc['filename']} ---\n\n"
        combined += doc['content']
    return combined.strip()

# Geriye dönük uyumluluk için eski fonksiyon adı
def get_pdf_text_legacy(pdf_docs):
    """Eski get_pdf_text fonksiyonu ile uyumluluk."""
    return get_combined_text(pdf_docs)
